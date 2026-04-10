# -*- coding: utf-8 -*-
"""Generate SDACS v6 report as Word docx"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Style
style = doc.styles['Normal']
style.font.size = Pt(10)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    return h

def add_tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t

def add_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

IMG_DIR = r'C:\Users\sun47\Desktop\swarm-drone-atc\.claude\worktrees\stupefied-hugle\v5_images'
IMG = {
    'g0':  'c6f3b3d148e074da06cbee1afedfc120699640e4.png',
    'g1':  '78cb17427a601bb078076de230d1279ef505df5a.png',
    'g2':  'e0271e0a50ac6bc78f0398f2c240f4e573bbe062.png',
    'g3':  'df06668eb85513a5ac7486fabc179561c38c8647.png',
    'g4':  '80746b62fb86c64a171c54eec1e860c582be58ce.png',
    'g5':  'cfa38868f013731afcd63c1435b362661dee6a83.png',
    'g6':  'ce7ff024213fe4ac9b79445fa88b3974a8322ea9.png',
    'g7':  '4894b105349392aeef769d0f24ad5b2d49e725a3.png',
    'g8':  '152b2021b23eb93a4da13a6a82826edb8f0e7e2c.png',
    'g9':  'f65473721ace1a8dabd5f3081c045a4e62b3dcfc.png',
    'g10': 'dcd74a92f4ea60cab7a9d3640ebfa33963603eb3.png',
    'g11': 'd0b4bce5521f952cd294984b8aa919b2eaaea0ea.png',
    'g12': 'ac244e51158df836c161d6734de59149585a941d.png',
    'g13': 'ca3c9b5e10bd3cd0f142ee2967693bf940d40e81.png',
    'g14': '7f6dcc79e78a06e547906f2d89ba69fc67bf5ad5.png',
}

def add_img(key, caption, width_cm=14):
    import os
    ipath = os.path.join(IMG_DIR, IMG[key])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(ipath, width=Cm(width_cm))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = c.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ═══ 표지 ═══
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SDACS')
r.font.size = Pt(36)
r.bold = True
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Swarm Drone Airspace Control Automation System')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\uad70\uc9d1\ub4dc\ub860 \uacf5\uc5ed\ud1b5\uc81c \uc790\ub3d9\ud654 \uc2dc\uc2a4\ud15c')
r.font.size = Pt(16)
r.bold = True

doc.add_paragraph()

for txt, sz in [
    ('\uce85\uc2a4\ud1a4\ub514\uc790\uc778 \xb7 \uc0ac\uc5c5 \ubc1c\ud45c \uc790\ub8cc', 12),
    ('\uad6d\ub9bd\ubaa9\ud3ec\ub300\ud559\uad50 \ub4dc\ub860\uae30\uacc4\uacf5\ud559\uacfc  |  \uc7a5\uc120\uc6b0  |  2026', 11),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.size = Pt(sz)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"드론 산업의 미래는 기체가 아니라 두뇌에 있습니다."')
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_img('g0', '[그림 0] SDACS 핵심 성과 지표 요약', 14)
doc.add_page_break()

# ═══ 목차 ═══
add_h('목   차', 1)
toc = [
    '1.  문제 제기 \u2014 하늘길이 막히고 있습니다',
    '2.  기존 시스템의 한계',
    '3.  SDACS 핵심 아이디어 \u2014 이동형 가상 레이더 돔',
    '4.  시스템 구조 \u2014 4계층 아키텍처',
    '5.  핵심 알고리즘 \u2014 어떻게 충돌을 막는가',
    '6.  연구 프레임워크 \u2014 왜 스타크래프트인가',
    '7.  탐지 \u2192 퇴각 자동화 파이프라인',
    '8.  관제 대시보드 및 JARVIS AI',
    '9.  성능 검증 결과',
    '10. 63개 시나리오 실측 테스트 결과 [신규]',
    '11. 광주시 테스트베드 전략',
    '12. 활용 분야 및 사회\xb7산업적 가치',
    '13. 개발 로드맵 (16주 + 장기 10년)',
    '14. 결론',
    '15. 특허 관계도',
    '16. 특허 출원 전략',
]
for item in toc:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ═══ 1. 문제 제기 ═══
add_h('1  문제 제기 \u2014 하늘길이 막히고 있습니다', 1)
add_img('g1', '[그림 1] 문제의 규모 — 4가지 핵심 수치', 14)
add_img('g2', '[그림 2] 기존 레이더의 한계 — 도심 저고도 67%가 사각지대', 14)
doc.add_paragraph(
    '지금 이 순간에도 전국의 하늘 위에서 수십만 대의 드론이 날아다니고 있습니다. '
    '택배를 배달하고, 논밭에 농약을 뿌리고, 건물을 점검하며, 2030년에는 도심 항공 모빌리티(UAM) 서비스도 본격화됩니다. '
    '전 세계 드론 시장은 연평균 25% 성장 중이며, 국내만 해도 등록 드론 수가 이미 90만 대를 돌파했습니다.'
)
doc.add_paragraph(
    '문제는 이 드론들이 모두 같은 저고도 하늘(지상 120m 이하)을 공유한다는 점입니다. '
    '기존 항공 관제 시스템은 대형 항공기 중심으로 설계되어 소형 드론은 레이더에 잡히지조차 않습니다. '
    '도심 저고도의 67%가 사각지대이며, 위험 상황에도 반응하는 데 평균 5분이 걸립니다.'
)
add_tbl(
    ['현재 상황', '수치', '의미'],
    [
        ['국내 드론 등록', '90만 대+', '매년 30% 이상 증가 중'],
        ['도심 저고도 사각지대', '67%', '기존 레이더가 탐지 못하는 구간'],
        ['수동 관제 반응 시간', '평균 5분', '고속 드론 위협에 즉각 대응 불가'],
        ['고정 레이더 구축 비용', '수억 원 + 6개월', '긴급 상황\xb7소규모 적용 불가'],
        ['2030년 드론 ATC 시장', '$15B 예상', '선제적 기술 확보가 시장 선점 열쇠'],
    ]
)

# ═══ 2. 기존 시스템의 한계 ═══
add_h('2  기존 시스템의 한계', 1)
doc.add_paragraph(
    '현재 운용 중인 드론 교통관리 시스템들은 각각 구조적 한계를 가지고 있습니다. '
    '성능이 부족한 것이 아니라, 설계 철학 자체가 새로운 드론 시대에 맞지 않습니다.'
)
add_tbl(
    ['시스템', '핵심 문제', 'SDACS 해결 방식'],
    [
        ['K-UTM', '중앙 서버 의존 \u2192 서버 다운 시 전체 마비', '분산 Mesh \u2192 드론 10% 고장 시에도 90% 작동'],
        ['NASA UTM', '고고도 최적화. 도심 저고도 적용 어려움', '군집드론을 공중에 배치해 사각지대 동적 커버'],
        ['고정형 레이더', '수억원+6개월. 도심 67% 미감시', '드론 10대로 30분 내 설치. 비용 90%+ 절감'],
        ['드론쇼 방식', '사전 경로만 실행. 돌발 대응 불가', '실시간 AI 자율 판단. 집단 지능 창발'],
    ]
)
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('드론쇼 vs SDACS: ')
r.bold = True
doc.add_paragraph('드론쇼: 중앙에서 설계한 계획을 각 드론이 실행하는 하향식', style='List Bullet')
doc.add_paragraph('SDACS: 단순 규칙을 따르는 드론들이 소통하며 집단 지능이 생겨나는 상향식', style='List Bullet')

# ═══ 3. 핵심 아이디어 ═══
add_h('3  SDACS 핵심 아이디어 \u2014 이동형 가상 레이더 돔', 1)
add_img('g3', '[그림 3] 이동형 가상 레이더 돔 — 20대 관제 드론이 Mesh 네트워크 형성', 14)
doc.add_paragraph(
    '"레이더를 땅에 설치하는 대신, 드론 자체가 레이더가 되면 어떨까?" \u2014 '
    'SDACS는 이 단순한 발상에서 출발했습니다. 20대의 관제 전담 드론이 공중에 올라가 '
    '그물망 형태의 감시 체계(가상 레이더 돔)를 스스로 구성합니다.'
)
add_h('[v6 보완] 실측 성능 데이터 업데이트', 2)
add_tbl(
    ['군집 규모', 'v5 해결률', 'v6 실측 해결률', '변화'],
    [
        ['50대', '97.9%', '99.9%', '+2.0%p'],
        ['100대', '98.9%', '99.5%', '+0.6%p'],
        ['200대', '70%', '98.7%', '+28.7%p 대폭 개선'],
        ['500대', '50% 이하', '97.4~98.9%', '+47%p 이상 혁신적'],
    ]
)
add_box('\u2192 500대 대규모 군집에서도 97%+ 안정적 운용 실증 완료')
doc.add_paragraph()
add_tbl(
    ['혁신 포인트', '설명'],
    [
        ['즉시 배치 (30분)', '기존 6개월 \u2192 30분 (99.7% 단축)'],
        ['5겹 입체 충돌 방지', '경로설정\u219290초 예측\u2192AI 자동 회피\u2192비상 브레이크\u2192안전 모드. 5중 안전망.'],
        ['극한 날씨 대응', '8가지 기상 시나리오. 63개\xd75회=316회 검증.'],
        ['[신규] GPU 가속 연산', 'WebGPU Compute Shader. 500대 실시간 연산.'],
    ]
)

# ═══ 4. 시스템 구조 ═══
add_h('4  시스템 구조 \u2014 4계층 아키텍처', 1)
add_img('g4', '[그림 4] SDACS 4계층 아키텍처', 14)
doc.add_paragraph('SDACS는 4개의 독립 계층이 각자 역할에 집중하면서 긴밀히 연결됩니다.')
add_tbl(
    ['계층', '역할', '핵심 기술', '비유'],
    [
        ['Layer 4 사용자', '인터페이스', '3D 관제, AI 비서, 모바일앱', '관제사 화면'],
        ['Layer 3 백엔드', '두뇌\xb7분석', '고속 통신, 타이머, MC 검증', '시스템의 두뇌'],
        ['Layer 2 지상', '관제소', '드론 OS, 장거리 통신, 공역 분할', '지휘 본부'],
        ['Layer 1 공중', '드론 군집', '비행 제어, 군집 알고리즘, APF, AI 카메라', '하늘의 드론들'],
    ]
)
add_h('[v6 보완] Layer 4 구현 기술 스택', 2)
for t in [
    'Three.js WebGL 3D 실시간 렌더링 (브라우저)',
    'WebGPU Compute Shader (WGSL) APF 병렬 연산',
    'Web Worker CPU 병렬 연산 (GPU 미지원 시 자동 전환)',
    'Spatial Hash O(N\xb7k) 이웃 탐색',
    'Canvas 2D 실시간 분석 차트 7종',
    '7대 광역시 도시환경 데이터 내장 (핫스왑)',
    '63개 시나리오 내장 선택 실행',
    'Dash/Plotly 3D 관제 대시보드',
]:
    doc.add_paragraph(t, style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('데이터 흐름: 드론(10Hz) \u2192 관제(1Hz) \u2192 서버 \u2192 관제사. 전체 1초 이내.')

# ═══ 5. 핵심 알고리즘 ═══
add_h('5  핵심 알고리즘 \u2014 어떻게 충돌을 막는가', 1)
add_img('g5', '[그림 5] 5겹 안전망 — 알고리즘 계층별 역할', 14)
add_img('g6', '[그림 6] 자동 우선순위 전환 — 상황에 따라 5단계로 자동 변경', 14)
add_h('5겹 안전망 (일상 비유)', 2)
add_tbl(
    ['층', '알고리즘', '일상 비유'],
    [
        ['1층', 'CBS 다중 경로 계획', '내비게이션: 출발 전 최적 경로 설정'],
        ['2층', 'CPA 충돌 예측 (90초 전)', '교통 레이더: 90초 전 부딪힘 예고'],
        ['3층', 'APF 자석형 회피 (GPU)', '같은 극 자석: 가까워지면 자동 밀어냄'],
        ['4층', '비상 브레이크', '급정거: 최후 순간 충돌 방지'],
        ['5층', '안전 모드 전환 [신규]', '비상구: 극한 상황 시 안전 착륙'],
    ]
)

add_h('자동 우선순위 전환', 2)
add_tbl(
    ['우선순위', '모드', '발동 조건', '자동 조치'],
    [
        ['P0', 'EMERGENCY', '충돌 임박/배터리 임계', '전체 비상 회피, 관제사 알림'],
        ['P1', 'DECONFLICT', '10m 이내 드론 감지', '자석형 힘 즉각 회피'],
        ['P2', 'MISSION', '임무 이착륙 중', '임무 경로 독점'],
        ['P3', 'CRUISE', '정상 비행', '경로 모니터링'],
        ['P4', 'IDLE', '위협 없음', '호버링 대기, 10초 스캔'],
    ]
)
add_box('[v6] APF 강풍 모드: 풍속 10m/s 초과 시 APF_PARAMS_WINDY 자동 전환 \u2192 안전 거리 확대')

# ═══ 6. 연구 프레임워크 ═══
add_h('6  연구 프레임워크 \u2014 왜 스타크래프트인가', 1)
add_img('g7', '[그림 7] 게임 AI → 실제 드론 ATC Sim-to-Real Transfer 프레임워크', 14)
add_img('g8', '[그림 8] 군집 규모별 제어 성능 — 권장 군집 크기: 50대', 14)
p = doc.add_paragraph()
r = p.add_run('핵심 연구 질문: ')
r.bold = True
p.add_run('"SC2에서 학습된 스웜 지능을 실제 드론 관제에 전이할 수 있는가?" \u2014 전 세계 미개척 영역.')
doc.add_paragraph('SC2 봇: 645단계, 404개 테스트, 797개 파일.')

add_h('[v6] 군집 규모별 성능 \u2014 실측 전면 업데이트', 2)
add_tbl(
    ['드론 수', '해결률', '병목', '해결책'],
    [
        ['10대', '100%', '없음', '기본 운용'],
        ['50대', '99.9% [\u219197.9%]', '없음', '권장 군집 크기'],
        ['100대', '99.5% [\u219198.9%]', '통신 대역폭', 'Edge Computing'],
        ['200대', '98.7% [\u219170%]', '연산량', 'GPU 가속 + 리더-팔로워'],
        ['500대', '97.4~98.9% [\u219150%]', '동기화', '로컬 분할 + WebGPU'],
    ]
)

# ═══ 7. 파이프라인 ═══
add_h('7  탐지 \u2192 퇴각 자동화 파이프라인 \u2014 1초 이내', 1)
add_img('g9', '[그림 9] 탐지 → 퇴각 자동화 파이프라인 (5단계 · 전체 1초 이내)', 14)
add_tbl(
    ['단계', '시간', '처리 내용', '핵심 기술'],
    [
        ['\u2460 탐지', '~50ms', '전파 탐지 + AI 카메라 30fps', '전파 + AI 카메라'],
        ['\u2461 식별', '즉시', '고유번호 DB 대조 \u2192 3분류', '고유번호 조회'],
        ['\u2462 타이머', '자동', '미등록 30초 카운트다운', '자동 타이머'],
        ['\u2463 경고', 'T-2분,T-0', 'SMS+앱 \u2192 최종(멀티채널)', '멀티채널 알림'],
        ['\u2464 퇴각', '에스컬', '포위 대형 + RF 경고', '군집 포위'],
    ]
)

# ═══ 8. 대시보드 ═══
add_h('8  관제 대시보드 및 JARVIS AI', 1)
add_tbl(
    ['패널', '기능'],
    [
        ['3D 레이더 맵', 'Three.js WebGL. 드론 상태별 색상. 60fps.'],
        ['AI 비서 (JARVIS)', '자연어 명령 즉시 실행.'],
        ['실시간 차트 [신규]', 'Canvas 2D 7종 분석 차트.'],
        ['이벤트 로그', '전 이력. 법적 근거 자료.'],
        ['센서 헬스', '1초마다 자동 점검.'],
        ['[신규] 성능 HUD', 'CPU/GPU/APF 연산 실시간 표시.'],
        ['[신규] 도시 선택', '7대 광역시 핫스왑.'],
        ['[신규] 시나리오 선택기', '63개 드롭다운 즉시 실행.'],
    ]
)

# ═══ 9. 성능 검증 ═══
add_h('9  성능 검증 결과', 1)
add_img('g10', '[그림 10] 기존 방식 vs SDACS — 성능 비교', 14)
add_h('[v6] 검증 규모 종합', 2)
for it in [
    'Monte Carlo: 384 조건 \xd7 100 시드 = 38,400가지',
    '63개 시나리오 \xd7 5회 = 316회 시뮬레이션 (2026-04-10)',
    '자동화 테스트: 2,668개 통과',
    'Python 3.10/3.11/3.12 CI/CD',
]:
    doc.add_paragraph(it, style='List Bullet')

doc.add_paragraph()
add_tbl(
    ['비교 항목', '기존', 'SDACS', '개선'],
    [
        ['준비 시간', '6개월', '30분', '99.7% 단축'],
        ['운영 인력', '5명(24h)', '1명', '80% 절감'],
        ['탐지 속도', '5분', '0.8초', '300배'],
        ['구축 비용', '수억 원+', '드론 10대', '90%+'],
        ['관제 대수', '20대', '500대 [신규]', '25배'],
        ['응답 지연', '200ms', '50ms', '4배'],
    ]
)

# ═══ 10. 시나리오 테스트 [전체 신규] ═══
doc.add_page_break()
add_h('10  63개 시나리오 실측 테스트 결과 [신규]', 1)
doc.add_paragraph('2026-04-10, 63개 시나리오 \xd7 5회 = 316회 실행. Three.js + WebGPU 환경.')

add_h('카테고리별 종합', 2)
add_tbl(
    ['카테고리', '수', '평균 해결률', '핵심 결과'],
    [
        ['기본', '5', '99.6%', '50~150대 안정 운용'],
        ['기상', '8', '92.4%', '극한기상 안전모드 작동'],
        ['도시환경(7광역시)', '21', '98.7%', '실제 도시 데이터 기반'],
        ['위협/침입', '4', '99.4%', '불법드론 대응 검증'],
        ['장애/복구', '5', '99.1%', '다중 장애에도 안정'],
        ['대규모/극한', '8', '96.3%', '500대 97%+'],
        ['특수 운용', '7', '99.5%', '편대/수색구조 등'],
        ['복합/고급', '5', '98.5%', '군사/재난 등'],
        ['전체 평균', '63', '96.9%', '55/63 \u2265 95%'],
    ]
)

add_h('7대 광역시 도시환경', 2)
add_tbl(
    ['도시', '협곡', '배달', '응급'],
    [
        ['서울', '99.0%', '98.9%', '98.5%'],
        ['부산', '98.3%', '98.7%', '99.0%'],
        ['인천', '99.0%', '98.7%', '98.4%'],
        ['대구', '98.7%', '98.7%', '99.1%'],
        ['광주', '98.9%', '98.6%', '98.2%'],
        ['대전', '98.4%', '98.6%', '98.9%'],
        ['울산', '98.9%', '98.8%', '98.4%'],
        ['평균', '98.7%', '98.7%', '98.6%'],
    ]
)

add_h('500대 대규모 군집', 2)
add_tbl(
    ['시나리오', '드론', '해결률', '특이사항'],
    [
        ['mega_swarm', '500', '97.4%', '최대 규모 기본'],
        ['mega_delivery', '500', '98.4%', '동시 택배 배송'],
        ['mega_storm', '500', '98.9%', '500대+폭풍'],
        ['total_war', '500', '98.5%', '복합 위협'],
        ['평균', '500', '98.3%', '500대도 98%+'],
    ]
)

add_h('충돌 0건 시나리오 (5개)', 2)
add_tbl(
    ['시나리오', '드론', '의미'],
    [
        ['formation_flight', '50', '편대 완벽 충돌 방지'],
        ['delivery_rush', '150', '대규모 배달에서도 0건'],
        ['microburst_storm', '150', '급변 기상 안전 모드'],
        ['extreme_stress', '200', '극한 스트레스 통과'],
        ['weather_hell', '200', '최악 기상 안전 착륙'],
    ]
)

# ═══ 11. 광주 테스트베드 ═══
add_h('11  광주시 테스트베드 전략', 1)
doc.add_paragraph('광주광역시(인구 150만, 501km\xb2)를 세계 최초 스웜 드론 ATC 실증 도시로.')
add_tbl(
    ['인프라', '규모', '활용'],
    [
        ['5G 기지국', '1,200+', '저지연 통신 (<10ms)'],
        ['CCTV', '5,000+', 'AI 비전 불법 드론 추적'],
        ['정책', 'AI 특별시+드론 샌드박스', '규제 샌드박스 가속'],
    ]
)
add_h('[v6] 광주 실측 데이터', 2)
add_tbl(
    ['시나리오', '드론', '해결률'],
    [
        ['gwangju_canyon', '150', '98.9%'],
        ['gwangju_delivery', '200', '98.6%'],
        ['gwangju_emergency', '120', '98.2%'],
    ]
)
add_box('\u2192 광주 3개 시나리오 모두 98%+. 실증 전 사전 검증 완료.')

# ═══ 12. 활용 분야 ═══
add_h('12  활용 분야 및 사회\xb7산업적 가치', 1)
add_tbl(
    ['분야', '활용 내용'],
    [
        ['도심 UAM\xb7드론 택배', '수천 대 안전 비행. 물류 기업 연계.'],
        ['국방\xb7안보', '수백 대 1인 통제. 적 드론 탐지.'],
        ['재난\xb7응급', '30분 만에 관제 구축. 자율 복구.'],
        ['환경\xb7사회', '탄소 감소. 산간\xb7도서 향상. 일자리 창출.'],
        ['[신규] 행사 보안', '축제\xb7경기장 임시 관제 즉시 배치.'],
        ['[신규] 농업\xb7감시', '광역 살포, 환경 모니터링, 산림 감시.'],
        ['[신규] 의료 수송', '응급 물자\xb7혈액\xb7장기 최우선 경로.'],
    ]
)

# ═══ 13. 로드맵 ═══
add_h('13  개발 로드맵', 1)
add_img('g11', '[그림 11] 개발 로드맵 — 측스톤 16주 · 장기 10년', 14)
add_tbl(
    ['Phase', '기간', '핵심 목표', '산출물'],
    [
        ['Phase 1 Sketch', '1~4주', 'SC2 봇 알고리즘 검증', 'SRS, SC2 프로토타입'],
        ['Phase 2 개념설계', '5~10주', '4계층 설계, 시뮬', '설계서, 시뮬레이터'],
        ['Phase 3 작품개발', '11~16주', '구현, AI 비서, 대시보드', '완성 관제 시스템'],
    ]
)
add_h('[v6] Phase 3 실제 달성', 2)
for a in [
    'Three.js 3D 시뮬레이터 완성',
    'WebGPU GPU 가속 연산 구현',
    '7대 광역시 도시환경 데이터 내장',
    '63개 시나리오 + 316회 실측',
    'Canvas 2D 분석 차트 7종',
    'Web Worker 병렬 연산',
    'Dash/Plotly 관제 대시보드',
    'JARVIS AI 비서 연동',
    'MC 38,400회 + 테스트 2,668개',
    'CI/CD GitHub Actions',
]:
    doc.add_paragraph(a, style='List Bullet')

doc.add_paragraph()
add_tbl(
    ['단계', '시기', '목표'],
    [
        ['단기', '2025~2026', 'Crazyflie 2~3대 야외 실비행'],
        ['중기', '2027~2028', '광주 실증. K-UTM 연동. IEEE. 특허 3건+.'],
        ['장기', '2029~2034', '광주 전역 상용화. 50개 도시 수출. SCI 5편+.'],
    ]
)

# ═══ 14. 결론 ═══
add_h('14  결론 \u2014 하늘길의 혁명', 1)
doc.add_paragraph(
    'SDACS는 드론 자체를 레이더로 삼아 AI가 자율 판단하고, '
    '일부 고장에도 전체 작동을 유지하는 분산형 자율 시스템입니다.'
)
add_h('[v6] 핵심 성과', 2)
for r_text in [
    '63개 시나리오 평균 충돌해결률 96.9%',
    '500대 군집에서도 97%+ 안정 운용',
    '7대 광역시 평균 98.7% (전 도시 98%+)',
    '극한 기상 안전 모드 자동 전환',
    'WebGPU GPU 가속 500대 실시간 연산',
    '동시 관제 100대 \u2192 500대 (5배)',
]:
    doc.add_paragraph(r_text, style='List Bullet')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"드론 산업의 미래는 기체가 아니라 두뇌에 있습니다."')
r.italic = True
r.bold = True
r.font.size = Pt(12)

# ═══ 15. 특허 관계도 ═══
add_h('15  특허 관계도', 1)
add_img('g12', '[그림 12] SDACS 특허 관계도 — 기존 특허와의 겹침 및 위험도 분석', 14)
add_tbl(
    ['기관', '겹침', '위험도', 'SDACS 차이'],
    [
        ['NASA', 'UTM 충돌 회피', '주의', '지상 vs 공중'],
        ['Amazon', 'Mesh+미인가 대응', '주의', '배달 부수 vs 전용 ATC'],
        ['NC주립대', '전파 탐지+고유번호', '주의', '지상탐지 vs 공중탐지'],
        ['Wing(Google)', '분산 경로, 충돌', '낮음', '서버 SW vs 물리 ATC'],
        ['ETRI', '자기조직화 Mesh', '낮음', '범용 vs ATC 특화'],
    ]
)

# ═══ 16. 특허 출원 ═══
add_h('16  특허 출원 전략', 1)
add_img('g13', '[그림 13] 특허 출원 가능 기술 5선 및 실행 절차 4단계', 14)
add_tbl(
    ['우선', '대상', '독창성', '종류'],
    [
        ['\u2605\u2605\u2605', '이동형 가상 레이더 돔', '세계 선례 없음', '시스템'],
        ['\u2605\u2605\u2605', 'SC2\u2192드론 전이', 'RTS\u2192실제 최초', '방법'],
        ['\u2605\u2605\u2606', '5단계 퇴각 파이프라인', 'DB 타이머 자동화', '방법/시스템'],
        ['\u2605\u2605\u2606', '바람 연동 자동 회피', '풍속\u2192안전거리', '방법'],
        ['\u2605\u2606\u2606', 'LLM 자연어 ATC', 'ATC+LLM 초기', '시스템'],
    ]
)
doc.add_paragraph()
add_tbl(
    ['시기', '단계', '할 일'],
    [
        ['지금', '1', '목포대 산학협력단 상담'],
        ['1~2주', '2', 'KIPRIS 조사 후 가출원'],
        ['출원 후', '3', '변리사 청구항 작성'],
        ['상용화', '4', 'PCT 국제 출원'],
    ]
)
p = doc.add_paragraph()
r = p.add_run('\u203b 발표/논문/GitHub 공개 전 반드시 가출원으로 날짜 확보!')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

# ═══ 변경 이력 ═══
doc.add_page_break()
add_h('v5 \u2192 v6 보완 사항 요약', 1)
add_img('g14', '[그림 14] 추가 콘텐츠 6가지 — 목적에 따라 선택 포함', 14)
add_h('수치 보정', 2)
for f in [
    '50대: 97.9%\u219199.9%, 100대: 98.9%\u219199.5%',
    '200대: 70%\u219198.7% (WebGPU)',
    '500대: 50%\u219197.4~98.9%',
    '관제 대수: 100대\u2192500대',
    '시나리오: 63개 316회 실측',
]:
    doc.add_paragraph(f, style='List Bullet')

add_h('신규/보강', 2)
for e in [
    '[10장] 63개 시나리오 실측 결과 (전체 신규)',
    'WebGPU GPU 가속 구현',
    'Layer 4 기술 스택 상세',
    '5겹 안전망 비유 + 5층 추가',
    'APF 강풍 모드 상세',
    '대시보드 패널 8개 상세',
    '검증 규모 종합',
    '광주 실측 데이터',
    '활용 분야 3개 추가',
    'Phase 3 달성 내용',
    '결론 수치 업데이트',
]:
    doc.add_paragraph(e, style='List Bullet')

# ═══ 저장 ═══
path = r'C:\Users\sun47\Desktop\SDACS_최종보고서_v6_이미지포함.docx'
doc.save(path)
print(f'Saved: {path}')
